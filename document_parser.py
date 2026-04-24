import re
import io
from dataclasses import dataclass
from typing import Dict, List, Optional
from enum import Enum


class SectionType(str, Enum):
    EXCELLENCE = "excellence"
    IMPACT = "impact"
    IMPLEMENTATION = "implementation"
    OPEN_SCIENCE = "open_science"
    GENDER_DIMENSION = "gender_dimension"
    DISSEMINATION = "dissemination"
    EXPLOITATION = "exploitation"
    WORK_PACKAGES = "work_packages"
    RISK_TABLE = "risk_table"
    DELIVERABLES = "deliverables"
    MILESTONES = "milestones"
    ETHICS = "ethics"
    CONSORTIUM = "consortium"
    ABSTRACT = "abstract"
    REFERENCES = "references"
    UNKNOWN = "unknown"


@dataclass
class ExtractedSection:
    section_type: SectionType
    title: str
    content: str
    page_start: int
    page_end: int
    word_count: int


@dataclass
class ParsedProposal:
    full_text: str
    sections: Dict[SectionType, ExtractedSection]
    total_pages: int
    total_words: int
    metadata: Dict[str, str]
    warnings: List[str]
    partner_names: List[str]
    person_names: List[str]
    trl_mentions: List[Dict]
    kpi_mentions: List[str]
    budget_figures: List[Dict]
    acronym: Optional[str]


SECTION_PATTERNS = {
    SectionType.EXCELLENCE: [
        r"(?i)^[\d\.]*\s*(?:section\s*)?1[\.\s]+excellence",
        r"(?i)^1\.\s+excellence",
        r"(?i)^[\d\.]*\s*excellence",
    ],
    SectionType.IMPACT: [
        r"(?i)^[\d\.]*\s*(?:section\s*)?2[\.\s]+impact",
        r"(?i)^2\.\s+impact",
        r"(?i)^[\d\.]*\s*impact",
    ],
    SectionType.IMPLEMENTATION: [
        r"(?i)^[\d\.]*\s*(?:section\s*)?3[\.\s]+(?:quality|implementation)",
        r"(?i)^3\.\s+(?:quality|implementation)",
        r"(?i)quality\s+and\s+efficiency\s+of\s+the\s+implementation",
    ],
    SectionType.OPEN_SCIENCE: [
        r"(?i)open\s+science\s+practices",
        r"(?i)open\s+science",
    ],
    SectionType.GENDER_DIMENSION: [
        r"(?i)gender\s+dimension",
        r"(?i)sex\s+and\s+gender",
    ],
    SectionType.DISSEMINATION: [
        r"(?i)dissemination\s+(?:and|&)\s+(?:exploitation|communication)",
        r"(?i)communication\s*,?\s+dissemination",
    ],
    SectionType.EXPLOITATION: [
        r"(?i)exploitation\s+(?:plan|strategy)",
        r"(?i)exploitation\s+of\s+results",
    ],
    SectionType.WORK_PACKAGES: [
        r"(?i)work\s+plan",
        r"(?i)work\s+packages?\s+(?:description|list)",
        r"(?i)^work\s+package\s+\d+",
    ],
    SectionType.RISK_TABLE: [
        r"(?i)risk\s+(?:management|assessment|table)",
        r"(?i)critical\s+risks",
    ],
    SectionType.DELIVERABLES: [
        r"(?i)(?:list\s+of\s+)?deliverables",
        r"(?i)^deliverable\s+\d+",
    ],
    SectionType.MILESTONES: [
        r"(?i)(?:list\s+of\s+)?milestones",
        r"(?i)^milestone\s+\d+",
    ],
    SectionType.ETHICS: [
        r"(?i)^ethics",
        r"(?i)ethics\s+self[-\s]?assessment",
    ],
    SectionType.CONSORTIUM: [
        r"(?i)consortium",
        r"(?i)participants?\s+(?:description|list)",
        r"(?i)capacity\s+of\s+participants",
    ],
    SectionType.ABSTRACT: [
        r"(?i)^abstract",
        r"(?i)^summary",
        r"(?i)^project\s+summary",
    ],
    SectionType.REFERENCES: [
        r"(?i)^references",
        r"(?i)^bibliography",
    ],
}


def _pdf_text(fb: bytes):
    """
    Safely extract text from PDF bytes using PyMuPDF.

    Important:
    page_count is captured before doc.close().
    This prevents: ValueError / RuntimeError: document closed
    """

    try:
        import fitz
    except ImportError as e:
        raise ImportError(
            "PyMuPDF paketi bulunamadı. requirements.txt içine PyMuPDF ekleyin."
        ) from e

    if not fb:
        raise ValueError("PDF dosyası boş okunuyor.")

    full_parts = []
    page_texts = []

    doc = fitz.open(stream=fb, filetype="pdf")

    try:
        page_count = len(doc)

        for i in range(page_count):
            page = doc.load_page(i)
            text = page.get_text("text") or ""

            page_no = i + 1
            page_texts.append((page_no, text))
            full_parts.append(f"\n--- PAGE {page_no} ---\n{text}")

    finally:
        doc.close()

    full_text = "\n".join(full_parts)

    if not full_text.strip():
        raise ValueError(
            "PDF metni çıkarılamadı. Dosya taranmış görsel PDF olabilir; OCR gerekir."
        )

    return full_text, page_count, page_texts


def _docx_text(fb: bytes):
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx paketi bulunamadı. requirements.txt içine python-docx ekleyin."
        ) from e

    if not fb:
        raise ValueError("DOCX dosyası boş okunuyor.")

    doc = Document(io.BytesIO(fb))

    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)

    if not full_text.strip():
        raise ValueError("DOCX metni çıkarılamadı veya belge boş.")

    pages = max(1, len(full_text.split()) // 350)

    return full_text, pages, [(1, full_text)]


def _detect_sections(full_text: str) -> Dict[SectionType, ExtractedSection]:
    sections: Dict[SectionType, ExtractedSection] = {}

    lines = full_text.split("\n")

    current_type = SectionType.UNKNOWN
    current_title = ""
    current_lines = []
    current_start_page = 1
    current_page = 1

    def flush_section(end_page: int):
        nonlocal current_type, current_title, current_lines, current_start_page

        if current_type == SectionType.UNKNOWN:
            return

        content = "\n".join(current_lines).strip()

        if not content:
            return

        sections[current_type] = ExtractedSection(
            section_type=current_type,
            title=current_title,
            content=content,
            page_start=current_start_page,
            page_end=end_page,
            word_count=len(content.split()),
        )

    for line in lines:
        stripped = line.strip()

        page_match = re.match(r"--- PAGE (\d+) ---", stripped)
        if page_match:
            current_page = int(page_match.group(1))
            continue

        if not stripped:
            current_lines.append("")
            continue

        matched_type = None

        for section_type, patterns in SECTION_PATTERNS.items():
            for pattern in patterns:
                if re.match(pattern, stripped):
                    matched_type = section_type
                    break
            if matched_type:
                break

        if matched_type:
            flush_section(current_page)

            current_type = matched_type
            current_title = stripped
            current_lines = []
            current_start_page = current_page
        else:
            current_lines.append(stripped)

    flush_section(current_page)

    return sections


def _extract_trl_mentions(full_text: str) -> List[Dict]:
    trl_mentions = []

    for match in re.finditer(
        r"(?i)TRL\s*[\-:]?\s*(\d)\s*(?:to|->|–|-)\s*(?:TRL\s*)?(\d)",
        full_text,
    ):
        trl_mentions.append(
            {
                "type": "range",
                "start_trl": int(match.group(1)),
                "end_trl": int(match.group(2)),
                "raw": match.group(0),
            }
        )

    if not trl_mentions:
        for match in re.finditer(r"(?i)TRL\s*[\-:]?\s*(\d)", full_text):
            trl_mentions.append(
                {
                    "type": "single",
                    "trl": int(match.group(1)),
                    "raw": match.group(0),
                }
            )

    return trl_mentions


def _extract_kpis(full_text: str) -> List[str]:
    kpis = []

    patterns = [
        r"(?i)(?:KPI|key\s+performance\s+indicator)s?\s*[:\-]\s*([^\n]+)",
        r"(?i)(?:indicator|metric)\s*[:\-]\s*([^\n]+)",
    ]

    for pattern in patterns:
        for match in re.finditer(pattern, full_text):
            value = match.group(0).strip()
            if value and value not in kpis:
                kpis.append(value)

    return kpis[:50]


def _extract_budget_figures(full_text: str) -> List[Dict]:
    figures = []

    patterns = [
        r"(?i)(?:€|EUR)\s*([\d,\.]+)\s*(?:million|m|M)?",
        r"(?i)([\d,\.]+)\s*(?:million|m|M)\s*(?:€|EUR)",
    ]

    for pattern in patterns:
        for match in re.finditer(pattern, full_text):
            figures.append(
                {
                    "raw": match.group(0),
                    "value": match.group(1),
                }
            )

    return figures[:100]


def _extract_partner_names(full_text: str) -> List[str]:
    patterns = [
        r"\b[A-Z][A-Za-z&\-\s]+(?:University|Institute|Institut|Centre|Center|Foundation|Association|GmbH|Ltd|AG|BV|Oy|SAS|SL|SRL)\b",
        r"\b(?:University|Institute|Institut|Centre|Center|Foundation|Association)\s+of\s+[A-Z][A-Za-z\-\s]+\b",
    ]

    partners = set()

    for pattern in patterns:
        for match in re.finditer(pattern, full_text):
            candidate = re.sub(r"\s+", " ", match.group(0)).strip()
            if 4 <= len(candidate) <= 120:
                partners.add(candidate)

    return sorted(partners)[:50]


def _extract_person_names(full_text: str) -> List[str]:
    persons = set()

    for match in re.finditer(
        r"(?:Prof\.|Dr\.|Mr\.|Ms\.|Mrs\.)\s+([A-Z][a-zA-Z\-]+\s+[A-Z][a-zA-Z\-]+)",
        full_text,
    ):
        persons.add(match.group(1).strip())

    return sorted(persons)[:50]


def _extract_acronym(full_text: str) -> Optional[str]:
    patterns = [
        r"(?i)(?:acronym)\s*[:\-]\s*([A-Z][A-Z0-9\-]{2,20})",
        r"(?i)(?:project\s+acronym)\s*[:\-]\s*([A-Z][A-Z0-9\-]{2,20})",
        r"(?i)(?:proposal\s+acronym)\s*[:\-]\s*([A-Z][A-Z0-9\-]{2,20})",
    ]

    for pattern in patterns:
        match = re.search(pattern, full_text)
        if match:
            return match.group(1)

    return None


def parse_proposal(file_bytes: bytes, filename: str) -> ParsedProposal:
    warnings = []

    if not filename:
        raise ValueError("Dosya adı alınamadı.")

    if not file_bytes:
        raise ValueError("Dosya boş okunuyor.")

    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        full_text, pages, page_texts = _pdf_text(file_bytes)

    elif lower_name.endswith((".docx", ".doc")):
        full_text, pages, page_texts = _docx_text(file_bytes)

    else:
        raise ValueError(f"Desteklenmeyen format: {filename}")

    total_words = len(full_text.split())
    sections = _detect_sections(full_text)

    for expected in [
        SectionType.EXCELLENCE,
        SectionType.IMPACT,
        SectionType.IMPLEMENTATION,
    ]:
        if expected not in sections:
            warnings.append(f"'{expected.value}' bölümü tespit edilemedi.")

    if total_words < 500:
        warnings.append(
            "Belge çok kısa görünüyor; tam Part B dosyası olmayabilir."
        )

    trl_mentions = _extract_trl_mentions(full_text)
    kpi_mentions = _extract_kpis(full_text)
    budget_figures = _extract_budget_figures(full_text)
    partner_names = _extract_partner_names(full_text)
    person_names = _extract_person_names(full_text)
    acronym = _extract_acronym(full_text)

    return ParsedProposal(
        full_text=full_text,
        sections=sections,
        total_pages=pages,
        total_words=total_words,
        metadata={
            "filename": filename,
            "parser": "document_parser_v2",
            "page_text_count": str(len(page_texts)),
        },
        warnings=warnings,
        partner_names=partner_names,
        person_names=person_names,
        trl_mentions=trl_mentions,
        kpi_mentions=kpi_mentions,
        budget_figures=budget_figures,
        acronym=acronym,
    )
